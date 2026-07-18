from flask import Flask, render_template, request, Response, jsonify
from openai import OpenAI
from dotenv import load_dotenv
import os
import sqlite3
from datetime import datetime
from docx import Document
from io import BytesIO
from rag import search_knowledge, load_knowledge

load_dotenv()

app = Flask(__name__)

client = OpenAI(
    api_key="sk-13a99e5a475f4105b3dd0edcebbd40fc",
    base_url="https://api.deepseek.com/v1"
)

# 初始化数据库
def init_db():
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic TEXT,
            doc_type TEXT,
            style TEXT,
            content TEXT,
            created_at TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# 启动时加载知识库
load_knowledge()

# ==================== 首页 ====================
@app.route('/')
def index():
    return render_template('index.html')


# ==================== 一步生成（带 RAG + 流式） ====================
@app.route('/generate', methods=['POST'])
def generate():
    topic = request.form.get('topic')
    doc_type = request.form.get('doc_type')
    style = request.form.get('style')

    # RAG：搜索知识库
    docs = search_knowledge(topic)
    reference = ""
    if docs:
        reference = "请参考以下公文范例：\n\n"
        for i, doc in enumerate(docs):
            reference += f"【参考案例 {i+1}】\n{doc['content'][:600]}\n\n"

    prompt = f"""
你是一位资深公文写作专家。请根据以下要求写一篇公文：

主题：{topic}
文种：{doc_type}
风格：{style}

{reference}

要求：
1. 格式规范，符合公文写作标准
2. 语言得体，符合所选风格
3. 结构清晰，有标题、正文、落款
    """

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一位资深的公文写作专家。"},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        stream=True
    )

    def generate_stream():
        full_content = ""
        for chunk in response:
            content = chunk.choices[0].delta.content
            if content:
                full_content += content
                yield f"data: {content}\n\n"

        conn = sqlite3.connect('history.db')
        c = conn.cursor()
        c.execute(
            "INSERT INTO records (topic, doc_type, style, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (topic, doc_type, style, full_content, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        conn.commit()
        conn.close()
        yield "data: [DONE]\n\n"

    return Response(generate_stream(), mimetype='text/event-stream')


# ==================== 生成大纲（两步模式第一步，不带 RAG 显示） ====================
@app.route('/outline', methods=['POST'])
def generate_outline():
    topic = request.form.get('topic')
    doc_type = request.form.get('doc_type')
    style = request.form.get('style')

    prompt = f"""
你是一位资深公文写作专家。请根据以下要求，先写一份大纲：

主题：{topic}
文种：{doc_type}
风格：{style}

要求：
1. 大纲要清晰，列出标题、正文要点、结尾
2. 每个要点用一句话概括
3. 只输出大纲，不要写全文
    """

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一位资深的公文写作专家。"},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        stream=True
    )

    def generate_stream():
        for chunk in response:
            content = chunk.choices[0].delta.content
            if content:
                yield f"data: {content}\n\n"
        yield "data: [DONE]\n\n"

    return Response(generate_stream(), mimetype='text/event-stream')


# ==================== 基于大纲生成全文（两步模式第二步，带 RAG） ====================
@app.route('/generate_full', methods=['POST'])
def generate_full():
    topic = request.form.get('topic')
    doc_type = request.form.get('doc_type')
    style = request.form.get('style')
    outline = request.form.get('outline')

    # RAG：搜索知识库
    docs = search_knowledge(topic)
    reference = ""
    if docs:
        reference = "请参考以下公文范例：\n\n"
        for i, doc in enumerate(docs):
            reference += f"【参考案例 {i+1}】\n{doc['content'][:600]}\n\n"

    prompt = f"""
你是一位资深公文写作专家。用户确认了以下大纲，请根据大纲扩写成一篇完整的公文。

主题：{topic}
文种：{doc_type}
风格：{style}

用户确认的大纲：
{outline}

{reference}

要求：
1. 严格按照大纲结构扩写
2. 语言得体，符合所选风格
3. 格式规范，有标题、正文、落款
    """

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一位资深的公文写作专家。"},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        stream=True
    )

    def generate_stream():
        full_content = ""
        for chunk in response:
            content = chunk.choices[0].delta.content
            if content:
                full_content += content
                yield f"data: {content}\n\n"

        conn = sqlite3.connect('history.db')
        c = conn.cursor()
        c.execute(
            "INSERT INTO records (topic, doc_type, style, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (topic, doc_type, style, full_content, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        conn.commit()
        conn.close()
        yield "data: [DONE]\n\n"

    return Response(generate_stream(), mimetype='text/event-stream')


# ==================== 历史记录相关 ====================
@app.route('/history', methods=['GET'])
def history():
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute("SELECT id, topic, doc_type, style, created_at FROM records ORDER BY id DESC")
    rows = c.fetchall()
    conn.close()

    history_list = []
    for row in rows:
        history_list.append({
            'id': row[0],
            'topic': row[1],
            'doc_type': row[2],
            'style': row[3],
            'created_at': row[4]
        })
    return jsonify(history_list)


@app.route('/delete/<int:record_id>', methods=['DELETE'])
def delete_record(record_id):
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute("DELETE FROM records WHERE id = ?", (record_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/record/<int:record_id>', methods=['GET'])
def get_record(record_id):
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute("SELECT content FROM records WHERE id = ?", (record_id,))
    row = c.fetchone()
    conn.close()
    if row:
        return jsonify({'content': row[0]})
    return jsonify({'content': ''})


@app.route('/export/<int:record_id>', methods=['GET'])
def export_word(record_id):
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute("SELECT topic, content FROM records WHERE id = ?", (record_id,))
    row = c.fetchone()
    conn.close()

    if not row:
        return "记录不存在", 404

    topic = row[0]
    content = row[1]

    doc = Document()
    doc.add_heading(topic, 0)
    doc.add_paragraph(content)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return Response(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': 'attachment; filename=doc.docx'
        }
    )


if __name__ == '__main__':
    app.run(debug=True)